import streamlit as st
import pandas as pd
from docx import Document
import os
import tempfile
from zipfile import ZipFile
import re
import random
import requests
import unicodedata
from collections import defaultdict
from difflib import SequenceMatcher

st.set_page_config(page_title="Générateur de Questionnaires", layout="wide")
st.title("Générateur de Questionnaires de Satisfaction à Chaud")

REQUIRED_COLS = ['nom', 'prénom', 'email', 'session', 'formation', 'formateur']


# =========================================================================
# Text normalization & matching
# =========================================================================

def normalize_text(s):
    """Lowercase, strip accents, normalize quotes and spaces for robust matching."""
    s = str(s).lower().strip()
    s = unicodedata.normalize('NFD', s)
    s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    s = s.replace('\u2019', "'").replace('\u2018', "'")
    s = s.replace('\u00a0', ' ')
    s = ' '.join(s.split())
    s = s.rstrip(':.,;?!')
    return s.strip()


# =========================================================================
# Paragraph / run manipulation helpers
# =========================================================================

def append_text_to_paragraph(para, value):
    """Append text as a new run, inheriting style from the last existing run."""
    value = str(value)
    if para.runs:
        tpl = para.runs[-1]
        new_run = para.add_run(value)
        new_run.font.bold = tpl.font.bold
        new_run.font.italic = tpl.font.italic
        new_run.font.underline = tpl.font.underline
        if tpl.font.size:
            new_run.font.size = tpl.font.size
        if tpl.font.name:
            new_run.font.name = tpl.font.name
        if tpl.font.color and tpl.font.color.rgb:
            new_run.font.color.rgb = tpl.font.color.rgb
    else:
        para.add_run(value)


def replace_checkbox_symbol(para, new_symbol):
    """Swap the first ☐/☑ in the paragraph, preserving run structure and styles."""
    for run in para.runs:
        if '☐' in run.text:
            run.text = run.text.replace('☐', new_symbol, 1)
            return True
        if '☑' in run.text:
            run.text = run.text.replace('☑', new_symbol, 1)
            return True
    return False


def remplacer_placeholders(paragraph, replacements):
    """Legacy placeholder replacement ({{nom}}, {{prenom}}, ...) for older templates."""
    if not paragraph.text or not any(k in paragraph.text for k in replacements):
        return
    full_text = ''.join(run.text for run in paragraph.runs)
    for key, value in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, str(value))
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font = first_run.font
        saved = {
            'bold': font.bold, 'italic': font.italic, 'underline': font.underline,
            'size': font.size, 'color': font.color.rgb if font.color else None,
            'name': font.name,
        }
    else:
        saved = None
    paragraph.clear()
    new_run = paragraph.add_run(full_text)
    if saved:
        new_run.font.bold = saved['bold']
        new_run.font.italic = saved['italic']
        new_run.font.underline = saved['underline']
        if saved['size']:
            new_run.font.size = saved['size']
        if saved['color']:
            new_run.font.color.rgb = saved['color']
        if saved['name']:
            new_run.font.name = saved['name']


# =========================================================================
# Section / label detection
# =========================================================================

IDENTITY_LABELS = [
    (['nom et prenom'], 'full_name'),
    (['adresse mail', 'adresse email', 'e-mail', 'email'], 'email'),
    (['reference de session', 'ref session'], 'session'),
    (['formateur'], 'formateur'),
]

SATISFACTION_QUESTION_KEYWORDS = [
    'qualite du service', 'qualite du contenu', 'pertinence du contenu',
    'clarte et organisation', 'qualite des supports', 'utilite des supports',
    'competence et professionnalisme', 'clarte des explications',
    'capacite', 'interactivite', 'globalement',
]


def detect_section(text_norm):
    if 'formation suivie' in text_norm:
        return 'formation'
    if 'handicap' in text_norm:
        return 'handicap'
    if any(kw in text_norm for kw in SATISFACTION_QUESTION_KEYWORDS):
        return 'satisfaction'
    return None


def detect_identity_label(text_norm):
    for labels, field in IDENTITY_LABELS:
        for label in labels:
            if text_norm.startswith(label) or text_norm == label:
                return field
    return None


# =========================================================================
# AI comment generation (OpenRouter)
# =========================================================================

def generer_commentaire_ia(openrouter_api_key, formation="la formation"):
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {openrouter_api_key}",
        "HTTP-Referer": "https://formation-entreprise.com",
        "X-Title": "Générateur Questionnaires",
        "Content-Type": "application/json"
    }

    prompt_points_forts = (
        f"""Tu viens de suivre une formation en {formation}. Rédige une seule phrase courte et naturelle pour décrire ton ressenti. N'utilise aucune liste, numéro, puce, tiret, ou mise en forme Markdown (comme **gras** ou *italique*). Ne commence pas par \"1.\", \"2.\", \"-\", ou autre symbole. Ne donne qu’une phrase, sans retour à la ligne et évite de commencer ta phrase par Je mais soit le plus aléatoire possible base toi sur les exemples que je t'ai donné ci dessous.

Inspire-toi librement des idées suivantes (mais sans les reprendre telles quelles) :
Explications claires et outils
Formation pratico-pratique
Plein de tips concrets
Le formateur est pédagogue
Cours interactifs accessibles à tout moment
Supports bien structurés
On repart avec un système clé en main
Résultats concrets dès la fin de la session

Ta réponse doit être 100 % humaine, spontanée et fluide. Une seule phrase. Pas de formatage. Pas de liste. Juste un ressenti authentique."""
    )

    prompt_remarques = (
        f"""Tu viens de terminer une formation en {formation}. Rédige une phrase spontanée, comme si tu laissais un commentaire libre à chaud sur ton ressenti général ou un point que tu voudrais partager, évite de commencer ta phrase par "Je" ou "J'ai" mais soit le plus aléatoire possible dans le début de la phrase.

Ta réponse peut contenir une appréciation générale, une suggestion, un ressenti, une émotion ou une remarque libre. Tu peux être personnel, tant que tu restes positif ou constructif.

⚠️ Important :
- Ne commence pas par un numéro, un tiret ou une liste
- Ne fais pas de mise en forme (pas de gras, italique, etc.)
- Ne retourne qu'une seule phrase courte, humaine, naturelle
- Ne commence pas tes phrases par un vrai coup de boost ou un vrai plaisir soit aléatoire dans tes propositions
Ta réponse doit être 100 % humaine, spontanée et fluide. Une seule phrase. Pas de formatage. Pas de liste. Juste un ressenti authentique."""
    )

    def appeler_api(prompt):
        data = {
            "model": "openai/gpt-4.1",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.5,
            "max_tokens": 100
        }
        response = requests.post(url, headers=headers, json=data, timeout=15)
        response.raise_for_status()
        raw = response.json()['choices'][0]['message']['content'].strip()
        options = [ligne.strip() for ligne in raw.splitlines() if ligne.strip()]
        return random.choice(options) if options else ""

    try:
        return appeler_api(prompt_points_forts), appeler_api(prompt_remarques)
    except Exception as e:
        st.error(f"Erreur API IA : {e}")
        return "", ""


# =========================================================================
# Main questionnaire generator
# =========================================================================

def generer_questionnaire(participant, template_path,
                          commentaire_points_forts='', commentaire_remarques=''):
    doc = Document(template_path)

    identity_values = {
        'full_name': f"{participant['prénom']} {participant['nom']}",
        'email': str(participant['email']),
        'session': str(participant['session']),
        'formateur': str(participant['formateur']),
    }

    # Legacy placeholders — harmless if absent in the new template
    legacy_replacements = {
        "{{nom}}": str(participant['nom']),
        "{{prenom}}": str(participant['prénom']),
        "{{email}}": str(participant['email']),
        "{{ref_session}}": str(participant['session']),
        "{{formation}}": str(participant['formation']),
        "{{formateur}}": str(participant['formateur']),
        "{{commentaire_points_forts}}": commentaire_points_forts or '',
        "{{commentaire_remarques}}": commentaire_remarques or '',
    }

    formation_target_norm = normalize_text(participant['formation'])

    # ---- PASS 1: classify paragraphs, fill identity labels, track sections ----
    classified = []
    current_section = None

    # Open-question slots (empty paragraphs right after the relevant question)
    points_forts_slot = None
    remarques_slot = None
    pending_slot_kind = None  # 'points_forts' | 'remarques' | None

    for para in doc.paragraphs:
        remplacer_placeholders(para, legacy_replacements)
        raw = para.text

        if not raw.strip():
            # Empty paragraph — if a slot was pending, this is it
            if pending_slot_kind == 'points_forts' and points_forts_slot is None:
                points_forts_slot = para
                pending_slot_kind = None
            elif pending_slot_kind == 'remarques' and remarques_slot is None:
                remarques_slot = para
                pending_slot_kind = None
            classified.append({'para': para, 'kind': 'empty'})
            continue

        is_checkbox_line = raw.lstrip().startswith(('☐', '☑'))
        text_norm = normalize_text(raw)

        if is_checkbox_line:
            option_text = raw.lstrip().lstrip('☐☑').strip()
            classified.append({
                'para': para, 'kind': 'checkbox',
                'section': current_section, 'option': option_text,
            })
        else:
            # Identity label → fill in place if empty
            field = detect_identity_label(text_norm)
            if field and ':' in raw:
                after_colon = raw.split(':', 1)[1].replace('\u00a0', ' ').strip()
                if not after_colon:
                    append_text_to_paragraph(para, identity_values[field])
            # Section detection
            section = detect_section(text_norm)
            if section:
                current_section = section
            # Open-question slot detection
            if 'points forts de cette formation' in text_norm:
                pending_slot_kind = 'points_forts'
            elif 'autres commentaires' in text_norm or 'commentaires / remarques' in text_norm:
                pending_slot_kind = 'remarques'
            elif 'points a ameliorer' in text_norm or 'points a am' in text_norm:
                pending_slot_kind = None  # don't auto-fill this one
            classified.append({'para': para, 'kind': 'question'})

    # ---- PASS 2: decide which checkboxes to tick ----
    section_options = defaultdict(list)
    for i, item in enumerate(classified):
        if item['kind'] == 'checkbox' and item['section']:
            section_options[item['section']].append((i, item['option']))

    checked_indices = set()

    # Formation: single best match across all options
    if 'formation' in section_options:
        best_i, best_score = None, 0.0
        for i, opt in section_options['formation']:
            opt_norm = normalize_text(opt)
            score = 1.0 if opt_norm == formation_target_norm \
                else SequenceMatcher(None, formation_target_norm, opt_norm).ratio()
            if score > best_score:
                best_score, best_i = score, i
        if best_i is not None and best_score >= 0.70:
            checked_indices.add(best_i)

    # Satisfaction: each contiguous group of ☐ lines = one question's options.
    # Random choice between "Très satisfait" / "Satisfait" per group.
    if 'satisfaction' in section_options:
        groups = []
        current_group = []
        last_i = None
        for i, opt in section_options['satisfaction']:
            if last_i is None or all(
                classified[j]['kind'] == 'empty' for j in range(last_i + 1, i)
            ) and (i - last_i) <= 3:
                current_group.append((i, opt))
            else:
                if current_group:
                    groups.append(current_group)
                current_group = [(i, opt)]
            last_i = i
        if current_group:
            groups.append(current_group)

        for group in groups:
            answer = random.choice(['Très satisfait', 'Satisfait'])
            answer_norm = normalize_text(answer)
            for i, opt in group:
                if normalize_text(opt) == answer_norm:
                    checked_indices.add(i)
                    break

    # Handicap: "Non concerné"
    if 'handicap' in section_options:
        for i, opt in section_options['handicap']:
            if normalize_text(opt) == 'non concerne':
                checked_indices.add(i)
                break

    # ---- PASS 3: apply ☐ / ☑ to every checkbox paragraph ----
    for i, item in enumerate(classified):
        if item['kind'] != 'checkbox':
            continue
        replace_checkbox_symbol(item['para'], '☑' if i in checked_indices else '☐')

    # ---- PASS 4: write AI comments into their slots (if generated) ----
    if commentaire_points_forts and points_forts_slot is not None:
        points_forts_slot.add_run(str(commentaire_points_forts))
    if commentaire_remarques and remarques_slot is not None:
        remarques_slot.add_run(str(commentaire_remarques))

    # Save
    safe_prenom = re.sub(r'[^a-zA-Z0-9]', '_', str(participant['prénom']))
    safe_nom = re.sub(r'[^a-zA-Z0-9]', '_', str(participant['nom']))
    filename = f"Questionnaire_satisfaction_{safe_prenom}_{safe_nom}_{participant['session']}.docx"
    output_path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(output_path)
    return output_path


# =========================================================================
# Streamlit UI
# =========================================================================

st.markdown("### Étape 1 : Importation des fichiers")
col1, col2 = st.columns(2)
with col1:
    excel_file = st.file_uploader("Fichier Excel des participants", type="xlsx")
with col2:
    template_file = st.file_uploader("Modèle Word", type="docx")

st.markdown("### Étape 2 : Configuration IA")
generer_ia = st.checkbox("Activer la génération de commentaires IA (nécessite clé API)")
openrouter_api_key = ""
frequence_ia = 1
if generer_ia:
    openrouter_api_key = st.text_input("Clé API OpenRouter", type="password")
    frequence_ia = st.slider("Fréquence de génération IA (1 sur x participants)",
                             min_value=1, max_value=10, value=4, step=1)

if excel_file and template_file:
    try:
        df = pd.read_excel(excel_file)

        # Normalize column names: strip whitespace + lowercase (handles 'Nom ', 'NOM', 'Prénom', etc.)
        df.columns = [str(c).strip().lower() for c in df.columns]

        missing = [c for c in REQUIRED_COLS if c not in df.columns]
        if missing:
            st.error(f"❌ Colonnes requises manquantes : {', '.join(missing)}")
            st.write("Colonnes détectées dans votre fichier :", list(df.columns))
            st.stop()

        # Drop rows where any required field is empty
        before = len(df)
        df = df.dropna(subset=REQUIRED_COLS).reset_index(drop=True)
        dropped = before - len(df)

        st.success(f"✅ {len(df)} participants détectés"
                   + (f" ({dropped} ligne(s) incomplètes ignorées)" if dropped else ""))

        # AI preview
        if generer_ia and openrouter_api_key:
            st.markdown("### 🎲 Prévisualiser un commentaire IA aléatoire")
            if st.button("🧠 Générer une prévisualisation pour un participant sélectionné aléatoirement"):
                candidats = df.iloc[::frequence_ia]
                if not candidats.empty:
                    participant_test = candidats.sample(1).iloc[0]
                    cmt_fort, cmt_libre = generer_commentaire_ia(
                        openrouter_api_key, participant_test['formation'])
                    st.markdown(f"**👤 Participant : {participant_test['prénom']} "
                                f"{participant_test['nom']} – Formation : {participant_test['formation']}**")
                    st.markdown("**🟢 Commentaire : Points forts**")
                    st.info(cmt_fort)
                    st.markdown("**💬 Commentaire : Remarque libre**")
                    st.info(cmt_libre)
                else:
                    st.warning("Aucun participant éligible avec la fréquence définie.")

        if st.button("Générer les questionnaires", type="primary"):
            with tempfile.TemporaryDirectory() as tmpdir:
                template_path = os.path.join(tmpdir, "template.docx")
                with open(template_path, "wb") as f:
                    template_file.seek(0)
                    f.write(template_file.read())

                zip_path = os.path.join(tmpdir, "Questionnaires.zip")

                with ZipFile(zip_path, 'w') as zipf:
                    progress_bar = st.progress(0)
                    errors = []
                    for idx, row in df.iterrows():
                        cmt_fort, cmt_libre = "", ""
                        if generer_ia and openrouter_api_key and idx % frequence_ia == 0:
                            cmt_fort, cmt_libre = generer_commentaire_ia(
                                openrouter_api_key, row['formation'])
                        try:
                            out = generer_questionnaire(row, template_path, cmt_fort, cmt_libre)
                            zipf.write(out, os.path.basename(out))
                        except Exception as e:
                            errors.append(f"{row['prénom']} {row['nom']}: {e}")
                        progress_bar.progress((idx + 1) / len(df),
                                              text=f"Progression : {idx+1}/{len(df)}")

                if errors:
                    st.warning(f"{len(errors)} erreur(s) :")
                    for e in errors[:10]:
                        st.text(f"  • {e}")

                with open(zip_path, "rb") as f:
                    st.balloons()
                    st.download_button(
                        "⬇️ Télécharger les questionnaires",
                        data=f,
                        file_name="Questionnaires_Satisfaction.zip",
                        mime="application/zip"
                    )
    except Exception as e:
        st.error(f"🚨 Erreur critique : {str(e)}")
